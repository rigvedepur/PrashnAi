import argparse
import os
import random
import re
from datetime import datetime

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def load_questions(input_path: str) -> list[dict]:
    """Parse questions from Energy_TossUps_All.txt created by extract_energy_tossups.py

    Expected block format per question:
      "<n>. <body>\nANSWER: <answer>\n\n"
    We will parse robustly via regex.
    """
    with open(input_path, 'r', encoding='utf-8') as f:
        text = f.read()

    # Split into entries: a line starting with number dot space
    # Keep the number to preserve original index if needed
    blocks = re.split(r"(?m)^(\d+)\.\s", text)
    # re.split returns [pre, num1, rest1, num2, rest2, ...]; discard preamble
    questions = []
    for i in range(1, len(blocks), 2):
        num = blocks[i]
        rest = blocks[i + 1]
        # The 'rest' contains body + possibly following blocks; stop at next double newline followed by number or end
        # More robust: find 'ANSWER:' within rest
        # Extract up to the next blank line that precedes another numbered question or end
        # First, try to split out the answer
        body, answer = split_body_answer(rest)
        body = fix_spacing_artifacts(body.strip())
        answer = fix_spacing_artifacts(answer.strip()) if answer else ""
        questions.append({
            'orig_number': int(num),
            'body': body,
            'answer': answer,
        })
    return questions


def split_body_answer(block_tail: str) -> tuple[str, str]:
    """Given the text right after '<n>. ', return (body, answer).
    We expect 'ANSWER:' line somewhere after body.
    """
    # Find ANSWER: ... to end of line
    m = re.search(r"(?im)^ANSWER\s*:\s*(.+)$", block_tail)
    if m:
        answer = m.group(1).strip()
        body = block_tail[:m.start()].strip()
        # Trim trailing blank lines from body
        body = re.sub(r"\s+$", "", body)
        return body, answer
    # Fallback: try inline ANSWER if on same paragraph
    m = re.search(r"(?is)ANSWER\s*:\s*(.+?)\s*(?:\n\s*\n|$)", block_tail)
    if m:
        answer = m.group(1).strip()
        body = block_tail[:m.start()].strip()
        return body, answer
    # No answer found; treat whole as body
    return block_tail.strip(), ""


def fix_spacing_artifacts(text: str) -> str:
    """Fix common PDF spacing artifacts conservatively.
    Mirrors the logic in extract_energy_tossups.py for consistency.
    """
    # Join hyphen followed by linebreaks/spaces
    text = re.sub(r"(\w)[ \t]*-[ \t]*\n[ \t]*(\w)", r"\1\2", text)
    # Also join hyphenations split across lines
    text = re.sub(r"(\w)-\s+(\w)", r"\1\2", text)
    # Remove spaces before punctuation
    text = re.sub(r"\s+([.,;:!?])", r"\1", text)
    # Normalize multiple spaces within lines
    text = "\n".join(re.sub(r"[ \t]{2,}", " ", ln) for ln in text.split("\n"))
    return text.strip()


def pick_unique_questions(questions: list[dict], k: int, seed: int | None = None) -> list[dict]:
    if seed is not None:
        rnd = random.Random(seed)
    else:
        rnd = random
    if k > len(questions):
        raise ValueError(f"Requested {k} questions but only {len(questions)} available")
    return rnd.sample(questions, k)


def build_pdf(output_path: str, selected: list[dict], title_suffix: str = "", start_number: int = 1) -> None:
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    # Tweak styles
    title_style = styles['Heading1']
    title_style.spaceAfter = 12
    q_style = ParagraphStyle(
        'Question', parent=styles['Normal'], fontSize=10.5, leading=14
    )
    a_style = ParagraphStyle('Answer', parent=styles['Normal'], fontSize=10, leading=13)

    story = []
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    title_text = f"Energy Toss-Up Set {title_suffix}".strip()
    story.append(Paragraph(title_text or "Energy Toss-Up Set", title_style))
    story.append(Paragraph(f"Generated: {now}", styles['Italic']))
    story.append(Spacer(1, 0.2 * inch))

    # Questions with inline answers
    for idx, q in enumerate(selected, start=start_number):
        # Body: prefix with Q{n}:
        body_html = escape_html(q['body']).replace('\n', '<br/>')
        q_para = Paragraph(f"Q{idx}: {body_html}", q_style)
        story.append(q_para)
        # Inline answer immediately after the question
        ans_text = q['answer'] or "[NOT FOUND]"
        a_para = Paragraph(f"Answer: {escape_html(ans_text)}", a_style)
        story.append(a_para)
        story.append(Spacer(1, 0.18 * inch))

    doc.build(story)


def build_docx_inline(output_path: str, selected: list[dict], title_suffix: str = "", start_number: int = 1) -> None:
    doc = Document()
    # Margins (approx 0.75")
    sections = doc.sections
    for sec in sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f"Energy Toss-Up Set {title_suffix}".strip() or "Energy Toss-Up Set")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    meta = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta.style = doc.styles['Intense Quote'] if 'Intense Quote' in doc.styles else None

    doc.add_paragraph("")

    for idx, q in enumerate(selected, start=start_number):
        p = doc.add_paragraph()
        run = p.add_run(f"Q{idx}: ")
        run.bold = True
        run.font.size = Pt(11)
        # Body with preserved line breaks
        for j, line in enumerate(q['body'].split('\n')):
            if j == 0:
                p.add_run(line)
            else:
                p = doc.add_paragraph(line)
        a = doc.add_paragraph()
        a_run = a.add_run(f"Answer: {q['answer'] or '[NOT FOUND]'}")
        a_run.font.size = Pt(10.5)
        doc.add_paragraph("")

    doc.save(output_path)


def build_docx_questions(output_path: str, selected: list[dict], title_suffix: str = "", start_number: int = 1) -> None:
    doc = Document()
    sections = doc.sections
    for sec in sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    run = title.add_run(f"Energy Toss-Up Questions {title_suffix}".strip() or "Energy Toss-Up Questions")
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    for idx, q in enumerate(selected, start=start_number):
        p = doc.add_paragraph()
        run = p.add_run(f"Q{idx}: ")
        run.bold = True
        run.font.size = Pt(11)
        for j, line in enumerate(q['body'].split('\n')):
            if j == 0:
                p.add_run(line)
            else:
                doc.add_paragraph(line)
        doc.add_paragraph("")

    doc.save(output_path)


def build_docx_answerkey(output_path: str, selected: list[dict], title_suffix: str = "", start_number: int = 1) -> None:
    doc = Document()
    sections = doc.sections
    for sec in sections:
        sec.top_margin = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    run = title.add_run(f"Energy Toss-Up Answer Key {title_suffix}".strip() or "Energy Toss-Up Answer Key")
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    for idx, q in enumerate(selected, start=start_number):
        p = doc.add_paragraph()
        run = p.add_run(f"Q{idx}: ")
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(q['answer'] or '[NOT FOUND]')

    doc.save(output_path)


def build_questions_pdf(output_path: str, selected: list[dict], title_suffix: str = "", start_number: int = 1) -> None:
    """Questions-only PDF (no answers)."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    title_style.spaceAfter = 12
    q_style = ParagraphStyle('Question', parent=styles['Normal'], fontSize=11, leading=15)

    story = []
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    title_text = f"Energy Toss-Up Questions {title_suffix}".strip()
    story.append(Paragraph(title_text or "Energy Toss-Up Questions", title_style))
    story.append(Paragraph(f"Generated: {now}", styles['Italic']))
    story.append(Spacer(1, 0.2 * inch))

    for idx, q in enumerate(selected, start=start_number):
        body_html = escape_html(q['body']).replace('\n', '<br/>')
        q_para = Paragraph(f"Q{idx}: {body_html}", q_style)
        story.append(q_para)
        story.append(Spacer(1, 0.2 * inch))

    doc.build(story)


def build_answerkey_pdf(output_path: str, selected: list[dict], title_suffix: str = "", start_number: int = 1) -> None:
    """Answer key PDF listing Qn and Answer only."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    title_style.spaceAfter = 12
    a_style = ParagraphStyle('Answer', parent=styles['Normal'], fontSize=11, leading=15)

    story = []
    now = datetime.now().strftime('%Y-%m-%d %H:%M')
    title_text = f"Energy Toss-Up Answer Key {title_suffix}".strip()
    story.append(Paragraph(title_text or "Energy Toss-Up Answer Key", title_style))
    story.append(Paragraph(f"Generated: {now}", styles['Italic']))
    story.append(Spacer(1, 0.2 * inch))

    for idx, q in enumerate(selected, start=start_number):
        ans_text = q['answer'] or "[NOT FOUND]"
        a_para = Paragraph(f"Q{idx}: {escape_html(ans_text)}", a_style)
        story.append(a_para)
        story.append(Spacer(1, 0.08 * inch))

    doc.build(story)


def escape_html(s: str) -> str:
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def main():
    parser = argparse.ArgumentParser(description="Generate printable PDFs with random Energy Toss-Up questions")
    parser.add_argument('--input', default='Energy_TossUps_All.txt', help='Path to Energy_TossUps_All.txt')
    parser.add_argument('--num', type=int, default=60, help='Number of questions per PDF')
    parser.add_argument('--num-pdfs', type=int, default=3, help='How many randomized outputs to generate')
    parser.add_argument('--seed', type=int, default=10, help='Base random seed for reproducibility')
    parser.add_argument('--output-dir', default='QUIZ_DIR', help='Directory to place generated PDFs')
    parser.add_argument('--start-number', type=int, default=1, help='Starting index for question numbering (for global numbering across PDFs)')
    parser.add_argument('--mode', choices=['inline', 'split'], default='split', help='inline: answers after each question; split: separate Questions and Answer Key PDFs')
    parser.add_argument('--format', choices=['pdf', 'docx', 'both'], default='docx', help='Output format')
    args = parser.parse_args()

    os.makedirs(args.output_dir, exist_ok=True)

    questions = load_questions(args.input)
    if len(questions) == 0:
        raise SystemExit("No questions parsed from input file.")
    if args.num > len(questions):
        raise SystemExit(f"Requested {args.num} questions but only {len(questions)} available")

    base_seed = args.seed
    for i in range(1, args.num_pdfs + 1):
        # Derive a different seed per output when a base seed is provided
        seed_i = None if base_seed is None else (base_seed + i * 9973)
        subset = pick_unique_questions(questions, args.num, seed=seed_i)
        # Give each output a unique name and title suffix
        suffix = f"#{i}"
        # Compute starting number per output if user wants to continue numbering across PDFs
        start_no = args.start_number if args.num_pdfs == 1 else (args.start_number + (i - 1) * args.num)

        if args.mode == 'inline':
            # Build inline format
            if args.format in ('pdf', 'both'):
                out_name = f"Energy_TossUps_Set_{i:02d}.pdf"
                out_path = os.path.join(args.output_dir, out_name)
                build_pdf(out_path, subset, title_suffix=suffix, start_number=start_no)
                print(f"Saved: {out_path}")
            if args.format in ('docx', 'both'):
                out_name = f"Energy_TossUps_Set_{i:02d}.docx"
                out_path = os.path.join(args.output_dir, out_name)
                build_docx_inline(out_path, subset, title_suffix=suffix, start_number=start_no)
                print(f"Saved: {out_path}")
        else:
            # Build split format
            if args.format in ('pdf', 'both'):
                q_name = f"Energy_TossUps_Set_{i:02d}_Questions.pdf"
                a_name = f"Energy_TossUps_Set_{i:02d}_AnswerKey.pdf"
                q_path = os.path.join(args.output_dir, q_name)
                a_path = os.path.join(args.output_dir, a_name)
                build_questions_pdf(q_path, subset, title_suffix=suffix, start_number=start_no)
                build_answerkey_pdf(a_path, subset, title_suffix=suffix, start_number=start_no)
                print(f"Saved: {q_path}")
                print(f"Saved: {a_path}")
            if args.format in ('docx', 'both'):
                q_name = f"Energy_TossUps_Set_{i:02d}_Questions.docx"
                a_name = f"Energy_TossUps_Set_{i:02d}_AnswerKey.docx"
                q_path = os.path.join(args.output_dir, q_name)
                a_path = os.path.join(args.output_dir, a_name)
                build_docx_questions(q_path, subset, title_suffix=suffix, start_number=start_no)
                build_docx_answerkey(a_path, subset, title_suffix=suffix, start_number=start_no)
                print(f"Saved: {q_path}")
                print(f"Saved: {a_path}")


if __name__ == '__main__':
    main()
